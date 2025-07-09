# Final Complete Code for streamlit app
# with OpenRouter API integration for insurance template processing
# support uploading docx template and multiple PDF reports
# Extract template fields, process PDFs with OCR if needed
# Use LLM to fill template with extracted values

import streamlit as st
import pandas as pd
import re
import requests
import json
from docx import Document
from docx.shared import Inches
import PyPDF2
import io
import zipfile
from typing import Dict, List, Optional
import os
from datetime import datetime
import tempfile
import fitz  # PyMuPDF
import easyocr
import cv2
import numpy as np
from pdf2image import convert_from_path
import warnings
warnings.filterwarnings("ignore")

# Global variable to store uploaded docx file
uploaded_docx_file = None

class InsuranceTemplateProcessor:
    def __init__(self, api_key: str = None):
        """
        Initialize the processor with OpenRouter API configuration
        """
        self.api_key = api_key
        self.api_url = "https://openrouter.ai/api/v1/chat/completions"
        self.headers = {
            "Authorization": f"Bearer {api_key}" if api_key else "",
            "Content-Type": "application/json"
        }
    
    def extract_template_keys(self, docx_file) -> List[str]:
        """
        Step 1: Extract all words inside square brackets [] from the template
        """
        try:
            doc = Document(docx_file)
            template_text = ""
            
            # Extract text from all paragraphs
            for paragraph in doc.paragraphs:
                template_text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        template_text += cell.text + "\n"
            
            # Find all text within square brackets
            keys = re.findall(r"\[(.*?)\]", template_text)
            
            # Remove duplicates and clean up
            unique_keys = list(set(keys))
            
            st.success(f"Found {len(unique_keys)} template fields: {unique_keys}")
            return unique_keys
            
        except Exception as e:
            # st.error(f"Error extracting template keys: {str(e)}")
            return []
    
    def extract_text_with_easyocr_pagewise(self, pdf_file, max_pages=35):
        """
        Extract text using EasyOCR with page-wise processing and progress tracking
        """
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(pdf_file.read())
            tmp_path = tmp.name

        try:
            reader = easyocr.Reader(['en'], gpu=False)
            full_text = ""

            doc = fitz.open(tmp_path)
            total_pages = min(len(doc), max_pages)

            for i in range(1, total_pages + 1):
                try:
                    # Convert PDF page to image
                    images = convert_from_path(tmp_path, dpi=150, first_page=i, last_page=i)
                    if not images:
                        continue

                    # Convert to format suitable for EasyOCR
                    img_array = np.array(images[0])
                    img_bgr = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
                    
                    # Perform OCR
                    result = reader.readtext(img_bgr)

                    # Extract text from OCR results
                    page_text = f"\n--- Page {i} ---\n"
                    for detection in result:
                        # Each detection is a tuple: (bbox, text, confidence)
                        if len(detection) >= 2 and detection[1].strip():
                            page_text += detection[1].strip() + "\n"
                    
                    full_text += page_text + "\n"

                except Exception as page_error:
                    continue

            # Clean up
            doc.close()
            os.unlink(tmp_path)
            
            return full_text

        except Exception as e:
            # Clean up on error
            try:
                doc.close()
                os.unlink(tmp_path)
            except:
                pass

            return ""
    
    def extract_pdf_text(self, pdf_files) -> str:
        """
        Step 2a: Extract text from multiple PDF files with intelligent method selection
        Files with 'Form' in name -> Use OCR
        Files with 'photo report' in name -> Use standard PDF reader
        Others -> Try standard first, then OCR fallback
        """
        combined_text = ""
        
        try:
            for pdf_file in pdf_files:
                file_name_lower = pdf_file.name.lower()
                
                # Determine extraction method based on filename
                if 'form' in file_name_lower:
                    try:
                        pdf_file.seek(0)
                        ocr_text = self.extract_text_with_easyocr_pagewise(pdf_file)
                        
                        if ocr_text.strip():
                            combined_text += f"\n--- Report from {pdf_file.name} (OCR Extraction) ---\n{ocr_text}\n"
                        else:
                            combined_text += f"\n--- Report from {pdf_file.name} (FAILED) ---\nOCR text extraction failed for this document.\n"
                            
                    except Exception as ocr_error:
                        combined_text += f"\n--- Report from {pdf_file.name} (FAILED) ---\nOCR text extraction failed for this document.\n"
                
                elif 'photo report' in file_name_lower:
                    try:
                        pdf_reader = PyPDF2.PdfReader(pdf_file)
                        pdf_text = ""
                        
                        for page in pdf_reader.pages:
                            pdf_text += page.extract_text() + "\n"
                        
                        if pdf_text.strip():
                            combined_text += f"\n--- Report from {pdf_file.name} (Standard Extraction) ---\n{pdf_text}\n"
                        else:
                            combined_text += f"\n--- Report from {pdf_file.name} (EMPTY) ---\nNo text found in this document.\n"
                            
                    except Exception as pdf_error:
                        combined_text += f"\n--- Report from {pdf_file.name} (FAILED) ---\nStandard text extraction failed for this document.\n"
                
            return combined_text
            
        except Exception as e:
            return ""
    
    def extract_key_values_with_llm(self, template_keys: List[str], pdf_text: str) -> Dict[str, str]:
        """
        Step 2b: Use LLM to extract key-value pairs from PDF text
        """
        try:
            # Create a prompt for the LLM
            prompt = f"""
            You are an expert insurance document processor. I have extracted text from insurance photo reports and need to fill out a template.

            Template fields to fill:
            {', '.join(template_keys)}

            Insurance report text:
            {pdf_text}

            Please extract the values for each template field from the insurance report text. Return the response as a JSON object where keys are the template fields and values are the extracted information.

            Rules:
            1. If a field cannot be found, use "Not Available" as the value
            2. Be precise and extract only relevant information
            3. For dates, use format: MM/DD/YYYY
            4. For monetary amounts, include currency symbol
            5. Return only the JSON object, no additional text

            Example format:
            {{
                "field1": "extracted_value1",
                "field2": "extracted_value2",
                "field3": "Not Available"
            }}
            """

            # Prepare the API request
            payload = {
                "model": "deepseek/deepseek-chat",  # Using free DeepSeek model
                "messages": [
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                "temperature": 0.3,
                "max_tokens": 2000
            }

            # Make API request
            response = requests.post(self.api_url, headers=self.headers, json=payload)
            
            if response.status_code == 200:
                result = response.json()
                llm_response = result['choices'][0]['message']['content']
                
                # Try to parse JSON from the response
                try:
                    # Clean the response to extract JSON
                    json_start = llm_response.find('{')
                    json_end = llm_response.rfind('}') + 1
                    json_str = llm_response[json_start:json_end]
                    
                    key_value_pairs = json.loads(json_str)
                    
                    return key_value_pairs
                    
                except json.JSONDecodeError:
                    # Fallback: create empty dict with template keys
                    return {key: "Not Available" for key in template_keys}
            else:
                return {key: "Not Available" for key in template_keys}
                
        except Exception as e:
            return {key: "Not Available" for key in template_keys}
    
    def fill_template(self, docx_file, key_value_pairs: Dict[str, str]) -> Document:
        """
        Step 3: Fill the template with extracted values
        """
        try:
            # Load the document
            doc = Document(docx_file)
            
            # Replace in paragraphs
            for paragraph in doc.paragraphs:
                for key, value in key_value_pairs.items():
                    if f'[{key}]' in paragraph.text:
                        paragraph.text = paragraph.text.replace(f'[{key}]', str(value))
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in key_value_pairs.items():
                            if f'[{key}]' in cell.text:
                                cell.text = cell.text.replace(f'[{key}]', str(value))
            return doc
            
        except Exception as e:
            return None

def main():
    st.set_page_config(
        page_title="Insurance Template Filler",
        page_icon="📋",
        layout="wide"
    )
    
    st.title("🏥 Insurance Template Filler Automation")
    
    # Sidebar for API configuration
    with st.sidebar:
        st.subheader("⚙️ Configuration")
        api_key = st.text_input(
            "OpenRouter API Key",
            type="password",
            value="YOUR-OPEN-ROUTER-API-KEY",
            help="Enter your OpenRouter API key. Leave empty to use demo mode."
        )
        
        st.subheader("📊 Processing Options")
        max_ocr_pages = st.slider(
            "Max OCR Pages per PDF",
            min_value=5,
            max_value=50,
            value=50,
            help="Limit OCR processing to prevent timeouts"
        )
        st.subheader("Template Format")
        st.markdown("""
        *Your .docx template should contain fields in square brackets:*
        - `[PolicyNumber]` - Will be filled with policy number
        - `[EstimatedCost]` - Will be filled with estimated cost
        """)
        st.subheader("Processing Features")
        st.markdown("""
        **1. Smart Text Extraction:**
        - 📄 Standard PDF text extraction (fast)
        - 🔍 OCR fallback
        
        **2. AI-Powered Extraction:**
        - 🤖 LLM interprets document content
        - 🎯 Automatic field mapping
        
        **3. Robust Processing:**
        - 🔄 Handles multiple file formats
        - ⚡ Optimized for performance
        """)

        with st.expander("🔧 System Requirements & Troubleshooting"):
            st.markdown("""
        **Required Python Packages:**
        ```
        streamlit, python-docx, PyPDF2, requests, pandas, easyocr, PyMuPDF, pdf2image, opencv-python, numpy, Pillow
        ```
        
        **Troubleshooting:**
        - **OCR not working**: Ensure poppler-utils is installed for pdf2image
        - **Memory issues**: Reduce max OCR pages in sidebar
        - **API errors**: Check your OpenRouter API key
        - **Template not filling**: Ensure fields are in [brackets] format
        
        **Performance Tips:**
        - Use standard PDF text when possible (faster than OCR)
        - Limit OCR pages for large documents
        - Optimize image quality in PDFs for better OCR results
        """)



    # Initialize processor
    processor = InsuranceTemplateProcessor(api_key)
    
    # Main interface
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📄 Upload Insurance Template")
        template_file = st.file_uploader(
            "Choose a .docx template file",
            type=['docx'],
            help="Upload your insurance template with fields in [brackets]"
        )
        uploaded_docx_file = template_file.name if template_file else None
        if template_file:
            
            # Extract template keys
            with st.spinner("Extracting template fields..."):
                template_keys = processor.extract_template_keys(template_file)
            
    
    with col2:
        st.header("📸 Upload Photo Reports")
        pdf_files = st.file_uploader(
            "Choose PDF photo reports",
            type=['pdf'],
            accept_multiple_files=True,
            help="Upload multiple PDF photo reports"
        )
        
        if pdf_files:
            st.success(f"✅ {len(pdf_files)} PDF files uploaded!")
    
    # Process button
    if template_file and pdf_files:
        if st.button("🚀 Process Documents", type="primary"):
            try:
                # Step 1: Extract template keys (already done above)
                if not template_keys:
                    st.error("No template keys found. Please check your template format.")
                    return
                
                # Step 2: Extract text from PDFs
                with st.spinner("Extracting text from PDF reports..."):
                    pdf_text = processor.extract_pdf_text(pdf_files)
                
                if not pdf_text:
                    st.error("Failed to extract text from PDF files.")
                    return
                
                # Show extracted text preview
                with st.expander("📄 View Extracted Text Preview"):
                    st.text_area("Extracted Text", pdf_text, height=200)
                
                # Step 3: Use LLM to extract key-value pairs
                with st.spinner("Using AI to extract information..."):
                    if api_key:
                        key_value_pairs = processor.extract_key_values_with_llm(template_keys, pdf_text)
                    else:
                        # Demo mode - simulate LLM response
                        st.info("Demo mode: Simulating LLM response")
                        key_value_pairs = {key: f"Demo_{key}_Value" for key in template_keys}
                
                # Display extracted key-value pairs
                st.subheader("🔍 Extracted Information")
                
                # Show key-value pairs in an expandable table
                with st.expander("📋 View Extracted Key-Value Pairs", expanded=False):
                    # Create DataFrame for better table display
                    df_data = {
                        "Field": list(key_value_pairs.keys()),
                        "Extracted Value": list(key_value_pairs.values())
                    }
                    df = pd.DataFrame(df_data)
                    
                    # Display as table
                    st.dataframe(
                        df,
                        use_container_width=True,
                        hide_index=True,
                        column_config={
                            "Field": st.column_config.TextColumn(
                                "Template Field",
                                width="medium",
                            ),
                            "Extracted Value": st.column_config.TextColumn(
                                "Extracted Value",
                                width="large",
                            )
                        }
                    )
                
                # Step 4: Fill template
                with st.spinner("Filling template..."):
                    filled_doc = processor.fill_template(template_file, key_value_pairs)
                
                if filled_doc:
                    # Save the filled document
                    output_buffer = io.BytesIO()
                    filled_doc.save(output_buffer)
                    output_buffer.seek(0)
                    
                    # Generate filename
                    output_filename = f"Filled_{uploaded_docx_file}"
                    
                    # Download button
                    st.download_button(
                        label="📥 Download Filled Template",
                        data=output_buffer.getvalue(),
                        file_name=output_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    # Show summary
                    st.subheader("📊 Processing Summary")
                    summary_data = {
                        "Metric": [
                            "Template Fields Found",
                            "PDF Reports Processed", 
                            "Fields Successfully Filled"
                        ],
                        "Value": [
                            len(template_keys),
                            len(pdf_files),
                            len([v for v in key_value_pairs.values() if v != "Not Available"])
                        ]
                    }
                    
                    st.dataframe(pd.DataFrame(summary_data))
                    
            except Exception as e:
                st.exception(e)  # Show full traceback for debugging
    
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; color: #666; padding: 20px;'>
        <p>🚀 <strong>Insurance Template Filler Automation Pipeline</strong></p>
        <p>Made with ❤️ by Arpit Patel</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()



    
